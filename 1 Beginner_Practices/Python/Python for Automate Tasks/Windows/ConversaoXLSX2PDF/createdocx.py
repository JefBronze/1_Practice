import docx;

from docx import Document

# Define the personalized text for each document
personalized_texts = [
    '1234', '5678', '9012'
]

# Create a new Word document
document = Document()

# Add the personalized content to the document
for text in personalized_texts:
    # Add the document header
    document.add_paragraph(f'Nº {text}\n\n')
    
    # Add the document body
    document.add_paragraph('JEFERSON BRONZE MOREIRA FILHO')
    document.add_paragraph('CPF: 051.734.899-36\n')
    document.add_paragraph('Endereço: Rua Arthur Suplicy de Lacerda, 249 – Seminário – Curitiba – Paraná – Brasil – CEP: 80740-210.')
    document.add_paragraph(f'\nDeclaro ter recebido, através de diversas operações em conta corrente, ao longo do mês de {text} referente ao pagamento de serviços especializados de mão de obra em construção civil (elétrica, predial, soldagem, alvenaria estrutural, hidráulica, acabamentos diversos, etc).')
    document.add_paragraph(f'\nCuritiba, {text}.\n\n')
    document.add_paragraph('Jeferson Bronze Moreira Filho')
    
    # Add a page break between each document
    document.add_page_break()

# Save the document to a specific location
document.save('C:/path/to/saved/document.docx')
